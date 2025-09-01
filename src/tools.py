from smolagents import Tool
from openai import OpenAI

import os

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

class Simple_Conversation_Tool(Tool):
    """
    A tiny small-talk tool that delegates response crafting to OpenAI.
    Great for greetings, pleasantries, and casual chit-chat.
    """
    name = "Simple_Conversation_Tool"
    description = (
        "Handles simple conversational messages like greetings and small talk. "
        "Use this tool when the user says hi/hello/hey/good morning, asks how you are, "
        "thanks you, or otherwise initiates casual chat."
    )

    inputs = {
        "message": {"type": "string", "description": "The user's chat message."}
    }

    output_type = "string"

    def __init__(self, api_key: str, model_id: str = "gpt-4o-mini"):
        super().__init__()
        self.client = OpenAI(api_key=api_key)
        self.model_id = model_id
        self.system_prompt = (
            "You are a friendly, concise customer support greeter. "
            "Respond in one or two sentences maximum. "
            "Be warm, helpful, and professional. "
            "If the user asks for support beyond greetings, invite them to share details."
        )

    def forward(self, message: str) -> str:
        # Lightweight guardrail: only handle small-talk; let the agent do other stuff
        # smalltalk_triggers = (
        #     "hi", "hello", "hey", "good morning", "good afternoon",
        #     "good evening", "how are you", "what's up", "sup", "salam",
        #     "assalam", "thanks", "thank you"
        # )
        msg_lower = message.lower()
        # if not any(t in msg_lower for t in smalltalk_triggers):
        #     return (
        #         "I'm here for quick greetings and small talk. "
        #         "Tell me a bit more about what you need help with!"
        #     )

        # Ask OpenAI to craft the short, friendly reply
        resp = self.client.chat.completions.create(
            model=self.model_id,
            messages=[
                {"role": "system", "content": self.system_prompt},
                {"role": "user", "content": message},
            ],
            temperature=0.6,
            max_tokens=80,
        )
        return resp.choices[0].message.content.strip()


from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from smolagents import Tool, ToolCallingAgent
from smolagents.models import OpenAIModel
from openai import OpenAI



class FAQTool(Tool):
    name = "FAQTool"
    description = (
        "Answer customer FAQ-style questions strictly using the local file "
        "`general_faqs.docx` (same directory as this script). "
        "Use for queries about shipping, returns, refunds, warranties, hours, policies, pricing, etc."
    )

    inputs = {
        "question": {"type": "string", "description": "User's FAQ-style question."}
    }

    output_type = "string"

    def __init__(
        self,
        api_key: str,
        model_id: str = "gpt-4o-mini",
        docx_path: str | Path = "general_faqs.docx",
        max_snippets: int = 8
    ):
        """
        - Loads the docx once and keeps paragraphs in memory.
        - Selects top-k relevant paragraphs to stay within token limits.
        """
        super().__init__()
        self.client = OpenAI(api_key=api_key)
        self.model_id = model_id
        self.docx_path = Path(docx_path)
        self.max_snippets = max_snippets
        self.paragraphs = self._load_docx_paragraphs(self.docx_path)

        self.system_prompt = (
            "You are a customer support assistant. Answer ONLY using the provided DATA. "
            "If the DATA does not contain the answer, say you don't have enough information. "
            "Be concise and actionable."
        )

    # ---- File loading helpers ----
    @staticmethod
    def _load_docx_paragraphs(path: Path) -> List[str]:
        if not path.exists():
            raise FileNotFoundError(f"FAQ file not found: {path.resolve()}")
        doc = Document(str(path))
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        # Merge tiny lines into bigger blocks (optional – keeps context coherent)
        merged: List[str] = []
        buf = []
        for p in paras:
            buf.append(p)
            if len(" ".join(buf)) > 300:  # simple merge threshold
                merged.append(" ".join(buf))
                buf = []
        if buf:
            merged.append(" ".join(buf))
        return merged

    # ---- Relevance selector: pick paragraphs most similar to the question ----
    @staticmethod
    def _keywordize(text: str) -> List[str]:
        tokens = re.findall(r"[a-z0-9]+", text.lower())
        stop = {
            "the","a","an","and","or","of","to","is","are","be","for","on","in",
            "at","by","with","from","as","that","this","it","we","you","your",
            "our","us","i"
        }
        return [t for t in tokens if t not in stop]

    def _score(self, paragraph: str, question: str) -> int:
        q = set(self._keywordize(question))
        p = set(self._keywordize(paragraph))
        return len(q & p)

    def _select_relevant_context(self, question: str) -> List[Tuple[int, str]]:
        scored = [(self._score(p, question), p) for p in self.paragraphs]
        scored.sort(key=lambda x: x[0], reverse=True)
        top = [(s, p) for s, p in scored[: self.max_snippets] if s > 0]
        # If nothing matched, still return a small sample to give the LLM something
        if not top and self.paragraphs:
            top = [(0, self.paragraphs[0])]
        # Attach indices to keep order stable (optional)
        return list(enumerate([p for _, p in top], start=1))

    # ---- Tool entrypoint ----
    def forward(self, question: str) -> str:
        # Lightweight trigger guard (optional; the agent usually routes to this tool)
        faq_cues = [
            "refund", "return", "shipping", "delivery", "warranty", "replacement",
            "order status", "track", "payment", "billing", "price", "discount",
            "hours", "open", "close", "support", "policy", "policies", "exchange",
            "cancel", "cancellation", "international", "customs", "duty", "tax"
        ]
        if not any(cue in question.lower() for cue in faq_cues):
            # Still proceed; the agent decided to call us. But we warn if it looks off.
            pass

        context_snippets = self._select_relevant_context(question)
        context_text = "\n\n".join(
            f"[Snippet {i}]\n{para}" for i, para in context_snippets
        )

        user_payload = (
            f"DATA:\n{context_text}\n\n"
            f"QUESTION:\n{question}\n\n"
            "INSTRUCTIONS:\n"
            "- Answer ONLY using DATA.\n"
            "- If not in DATA, say: 'I’m sorry, I don’t have that information.'\n"
            "- Keep it concise (2–5 sentences)."
        )

        resp = self.client.chat.completions.create(
            model=self.model_id,
            messages=[
                {"role": "system", "content": self.system_prompt},
                {"role": "user", "content": user_payload},
            ],
            temperature=0.2,
            max_tokens=300,
        )
        return resp.choices[0].message.content.strip()


class VerifyAnswerTool(Tool):
    """
    Verifies if the user's question has been sufficiently answered
    by the latest tool output. If yes, returns 'FINALIZE: <answer>'.
    Otherwise returns 'CONTINUE: <reason>'.
    """
    name = "VerifyAnswerTool"
    description = (
        "Given the original user question and the latest candidate answer, "
        "decide if the question has been fully answered. If sufficient, return "
        "'FINALIZE: <answer>'. If not, return 'CONTINUE: <brief reason>'."
    )

    inputs = {
        "question": {"type": "string", "description": "Original user question."},
        "candidate_answer": {"type": "string", "description": "Latest tool/model output you're evaluating."}
    }

    output_type = "string"

    def __init__(self, api_key: str, model_id: str = "gpt-4o-mini"):
        super().__init__()
        self.client = OpenAI(api_key=api_key)
        self.model_id = model_id
        self.system_prompt = (
            "You are a strict verifier. Determine if the candidate answer is correct "
            "answers the user's question. If YES, respond exactly as:\n"
            "FINALIZE: <Answer as is>\n"
            "If NO, respond exactly as:\n"
            "CONTINUE: <one-sentence reason what is missing or unclear>\n"
            "Do not add any other text."
        )

    def forward(self, question: str, candidate_answer: str) -> str:
        msg = (
            f"QUESTION:\n{question}\n\n"
            f"CANDIDATE_ANSWER:\n{candidate_answer}\n\n"
            "Does CANDIDATE_ANSWER fully answer QUESTION?"
        )
        resp = self.client.chat.completions.create(
            model=self.model_id,
            messages=[
                {"role": "system", "content": self.system_prompt},
                {"role": "user", "content": msg},
            ],
            temperature=0.0,
            max_tokens=120,
        )
        return resp.choices[0].message.content.strip()
    


